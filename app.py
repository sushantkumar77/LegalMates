# import streamlit as st
# import docx
# import os
# import io
# import re
# from dotenv import load_dotenv
# from openai import OpenAI  # We use the OpenAI library to call Cohere

# # --- Page Configuration ---
# st.set_page_config(
#     page_title="LegalEase AI (Smart Chat)",
#     page_icon="✍️",
#     layout="wide"
# )

# # --- Environment Variable & API Key ---
# load_dotenv()
# COHERE_API_KEY = os.getenv("COHERE_API_KEY")

# # Check for API key
# if not COHERE_API_KEY:
#     st.error("🚨 COHERE_API_KEY not found. Please set it in your .env or Streamlit secrets.")
#     st.stop()

# try:
#     # --- CONFIGURE OPENAI CLIENT FOR COHERE ---
#     # This setup points the standard OpenAI library to Cohere's API
#     client = OpenAI(
#         api_key=COHERE_API_KEY,
#         # This is the CORRECT URL to fix the 405 error
#         base_url="https://api.cohere.ai/compatibility/v1" 
#     )
#     COHERE_MODEL = "command-r-plus-08-2024" # <-- Use the new flagship model  # Use a powerful Cohere model
#     # --- END CLIENT CONFIGURATION ---
# except Exception as e:
#     st.error(f"Failed to configure Cohere-compatible client: {e}")
#     st.stop()

# # --- Helper Functions ---

# def extract_text_and_placeholders(file_bytes):
#     """
#     Extracts text and unique placeholders from a .docx file.
#     This regex is much more robust and covers most common syntaxes.
#     """
#     try:
#         doc = docx.Document(io.BytesIO(file_bytes))
#         full_text = []
#         for para in doc.paragraphs:
#             full_text.append(para.text)
#         for table in doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     for para in cell.paragraphs:
#                         full_text.append(para.text)
        
#         full_text_str = "\n".join(full_text)
        
#         # --- ROBUST REGEX ---
#         # Covers: {placeholder}, {{placeholder}}, [placeholder], <placeholder>, %placeholder%, __placeholder__, $placeholder
#         placeholder_regex = r"\{{1,2}.*?\}{1,2}|\[.*?\]|<.*?>|%.*?%|__.*?__|\$[a-zA-Z0-9_]+"
        
#         placeholders = list(set(re.findall(placeholder_regex, full_text_str))) # Unique list
        
#         return full_text_str, placeholders
#     except Exception as e:
#         st.error(f"Error reading .docx file: {e}")
#         return None, []

# def replace_placeholders_in_doc(file_bytes, replacements):
#     """
#     Replaces placeholders in a docx file (in memory) and returns the new file bytes.
#     This function replaces text while attempting to preserve formatting by operating on runs.
#     """
#     try:
#         doc = docx.Document(io.BytesIO(file_bytes))
#         for p in doc.paragraphs:
#             for key, value in replacements.items():
#                 if key in p.text:
#                     for run in p.runs:
#                         if key in run.text:
#                             run.text = run.text.replace(key, str(value))
#         for table in doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     for p in cell.paragraphs:
#                         for key, value in replacements.items():
#                             if key in p.text:
#                                 for run in p.runs:
#                                     if key in run.text:
#                                         run.text = run.text.replace(key, str(value))
#         file_stream = io.BytesIO()
#         doc.save(file_stream)
#         file_stream.seek(0)
#         return file_stream.getvalue()
#     except Exception as e:
#         st.error(f"Error replacing placeholders: {e}")
#         return None

# def clear_session_state_on_upload():
#     """Resets the session state when a new file is uploaded."""
#     keys_to_clear = [
#         "messages", "placeholders", "filled_values", 
#         "current_placeholder_index", "original_doc_bytes", 
#         "original_text", "api_history"
#     ]
#     for key in keys_to_clear:
#         if key in st.session_state:
#             del st.session_state[key]
    
#     # Re-initialize with the smart system prompt
#     st.session_state.messages = []
#     st.session_state.placeholders = []
#     st.session_state.filled_values = {}
#     st.session_state.current_placeholder_index = 0
#     st.session_state.original_doc_bytes = None  # <-- FIX
#     st.session_state.original_text = ""          # <-- FIX
#     st.session_state.api_history = [
#         {
#             "role": "system",
#             "content": """
#             You are a helpful and friendly assistant named 'LegalEase AI'. 
#             Your goal is to help a user fill in a document. 
#             I will give you placeholders one by one, like '{ClientName}' or '[DocumentDate]'.
#             Your job is to ask the user for the information to fill these placeholders.
            
#             RULES:
#             1.  Be conversational and natural. For example, instead of "What is {ClientName}?", ask "Who is the client for this agreement?" or "What's the client's full name?".
#             2.  Ask for only ONE piece of information at a time.
#             3.  When the user answers, confirm briefly (e.g., "Got it.", "Perfect.") and then immediately ask the question for the *next* placeholder I give you.
#             4.  Keep your questions clear and concise.
#             """
#         }
#     ]

# # --- Session State Initialization ---
# # This block now explicitly initializes all keys on the first run.
# if "messages" not in st.session_state:
#     st.session_state.messages = []
#     st.session_state.placeholders = []
#     st.session_state.filled_values = {}
#     st.session_state.current_placeholder_index = 0
#     st.session_state.original_doc_bytes = None  # <-- FIX
#     st.session_state.original_text = ""          # <-- FIX
#     st.session_state.api_history = [
#         {
#             "role": "system",
#             "content": """
#             You are a helpful and friendly assistant named 'LegalEase AI'. 
#             Your goal is to help a user fill in a document. 
#             I will give you placeholders one by one, like '{ClientName}' or '[DocumentDate]'.
#             Your job is to ask the user for the information to fill these placeholders.
            
#             RULES:
#             1.  Be conversational and natural. For example, instead of "What is {ClientName}?", ask "Who is the client for this agreement?" or "What's the client's full name?".
#             2.  Ask for only ONE piece of information at a time.
#             3.  When the user answers, confirm briefly (e.g., "Got it.", "Perfect.") and then immediately ask the question for the *next* placeholder I give you.
#             4.  Keep your questions clear and concise.
#             """
#         }
#     ]


# # --- Main App UI ---
# st.title("✍️ LegalEase AI: Conversational Document Filler")
# st.markdown("Upload your `.docx` template, and I'll help you fill in the blanks with a smart, conversational AI.")

# col1, col2 = st.columns([1, 1], gap="large")

# # --- Column 1: Upload & Chat ---
# with col1:
#     st.header("1. Upload & Fill")
    
#     uploaded_file = st.file_uploader(
#         "Upload your .docx template", 
#         type=["docx"],
#         on_change=clear_session_state_on_upload # Reset session on new file
#     )

#     if uploaded_file is not None and st.session_state.original_doc_bytes is None:
#         with st.spinner("Analyzing document..."):
#             file_bytes = uploaded_file.getvalue()
#             st.session_state.original_doc_bytes = file_bytes
#             text, placeholders = extract_text_and_placeholders(file_bytes)
            
#             if text is None:
#                 st.session_state.original_doc_bytes = None
#             elif not placeholders:
#                 st.warning("📄 No placeholders (like {Name} or [Date]) found.")
#                 st.session_state.original_doc_bytes = None
#             else:
#                 st.session_state.original_text = text
#                 st.session_state.placeholders = placeholders
#                 st.success(f"Found {len(placeholders)} placeholders!")
                
#                 with st.expander("Click to see all found placeholders"):
#                     st.write(placeholders)
                
#                 # --- SMART KICK-OFF PROMPT ---
#                 first_ph = st.session_state.placeholders[0]
#                 # This prompt tells the AI (which has the system prompt) to start the job.
#                 prompt_to_ai = f"Hello! I've uploaded a document. Here is the first placeholder: '{first_ph}'. Please ask me the first question."
                
#                 try:
#                     # Add our instruction to the AI's history
#                     st.session_state.api_history.append({"role": "user", "content": prompt_to_ai})
                    
#                     response = client.chat.completions.create(
#                         messages=st.session_state.api_history,
#                         model=COHERE_MODEL
#                     )
                    
#                     response_text = response.choices[0].message.content
                    
#                     # Add AI's response to both histories
#                     st.session_state.api_history.append({"role": "assistant", "content": response_text})
#                     st.session_state.messages.append({"role": "assistant", "content": response_text})
                    
#                 except Exception as e:
#                     st.error(f"Error with Cohere API: {e}")

#     # --- Conversational Chat ---
#     if st.session_state.original_doc_bytes is not None and st.session_state.placeholders:
#         st.markdown("---")
#         st.subheader("💬 Chat to Fill")

#         # Display chat history
#         for message in st.session_state.messages:
#             with st.chat_message(message["role"]):
#                 st.markdown(message["content"])

#         # Chat input
#         if prompt := st.chat_input("Your answer..."):
#             # Add user's answer to both histories
#             st.session_state.messages.append({"role": "user", "content": prompt})
#             st.session_state.api_history.append({"role": "user", "content": prompt})

#             with st.chat_message("assistant"):
#                 with st.spinner("Thinking..."):
#                     try:
#                         # Store the value
#                         current_index = st.session_state.current_placeholder_index
#                         current_ph = st.session_state.placeholders[current_index]
#                         st.session_state.filled_values[current_ph] = prompt
                        
#                         # Move to the next placeholder
#                         st.session_state.current_placeholder_index += 1
#                         next_index = st.session_state.current_placeholder_index
                        
#                         if next_index < len(st.session_state.placeholders):
#                             # --- SMART NEXT-QUESTION PROMPT ---
#                             next_ph = st.session_state.placeholders[next_index]
#                             ai_prompt = f"The next placeholder is: '{next_ph}'. Please ask me the question for this one."
#                         else:
#                             # --- SMART FINAL PROMPT ---
#                             ai_prompt = "That was the last placeholder! Please provide a brief, friendly message letting me know I'm all done and can review the document on the right."

#                         # Add our new instruction to the AI's history
#                         st.session_state.api_history.append({"role": "user", "content": ai_prompt})
                        
#                         response = client.chat.completions.create(
#                             messages=st.session_state.api_history,
#                             model=COHERE_MODEL
#                         )
                        
#                         response_text = response.choices[0].message.content
                        
#                         # Add AI's response to both histories
#                         st.session_state.api_history.append({"role": "assistant", "content": response_text})
#                         st.session_state.messages.append({"role": "assistant", "content": response_text})
                        
#                         st.markdown(response_text)
                    
#                     except Exception as e:
#                         st.error(f"Error with Cohere API: {e}")
#                         st.session_state.current_placeholder_index -= 1 # Roll back on error

# # --- Column 2: Review & Download (No changes needed) ---
# with col2:
#     st.header("2. Review & Download")
    
#     if st.session_state.original_doc_bytes is None:
#         st.info("Upload a document on the left to see a preview here.")
#     else:
#         with st.container(height=500, border=True):
#             st.subheader("Live Preview")
#             preview_text = st.session_state.original_text
            
#             for ph, val in st.session_state.filled_values.items():
#                 preview_text = preview_text.replace(ph, f"**{val}**") # Bold filled values
            
#             for ph in st.session_state.placeholders:
#                 if ph not in st.session_state.filled_values:
#                     preview_text = preview_text.replace(ph, f"_{ph}_") # Italicize unfilled

#             st.markdown(preview_text)
        
#         st.markdown("---")

#         all_filled = len(st.session_state.filled_values) == len(st.session_state.placeholders)
        
#         if all_filled and st.session_state.placeholders:
#             st.success("All fields filled! 🎉")
            
#             with st.spinner("Generating final document..."):
#                 final_doc_bytes = replace_placeholders_in_doc(
#                     st.session_state.original_doc_bytes,
#                     st.session_state.filled_values
#                 )
            
#             if final_doc_bytes:
#                 st.download_button(
#                     label="⬇️ Download Completed Document",
#                     data=final_doc_bytes,
#                     file_name=f"completed_{uploaded_file.name}",
#                     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#                     use_container_width=True
#                 )
#         elif st.session_state.placeholders:
#             progress = len(st.session_state.filled_values) / len(st.session_state.placeholders)
#             st.progress(progress, text=f"{len(st.session_state.filled_values)} / {len(st.session_state.placeholders)} fields filled")


import streamlit as st
import docx
import os
import io
import re
from dotenv import load_dotenv
from openai import OpenAI  # We use the OpenAI library to call Cohere

# --- 1. Page Configuration ---
st.set_page_config(
    page_title="LegalEase AI (Pro)",
    page_icon="✍️",
    layout="wide"
)

# --- 2. Environment Variable & API Client ---
load_dotenv()
COHERE_API_KEY = os.getenv("COHERE_API_KEY")

if not COHERE_API_KEY:
    st.error("🚨 COHERE_API_KEY not found. Please set it in your .env or Streamlit secrets.")
    st.stop()

try:
    client = OpenAI(
        api_key=COHERE_API_KEY,
        base_url="https://api.cohere.ai/compatibility/v1"
    )
    # FIX: Use the new, versioned model name
    COHERE_MODEL = "command-r-plus-08-2024"
except Exception as e:
    st.error(f"Failed to configure Cohere-compatible client: {e}")
    st.stop()

# --- 3. System Prompt (Upgraded for Smarter AI) ---
SYSTEM_PROMPT = """
You are a helpful and friendly legal assistant named 'LegalEase AI'.
Your goal is to help a user fill in a document. 
I will give you a placeholder (like '{ClientName}') AND its surrounding context from the document.
Your job is to use this context to ask a smart, relevant question to get the information.

RULES:
1.  **Use Context:** Don't just ask "What is {ClientName}?". If the context is "This agreement is between {ClientName} and ACME Corp...", ask "Who is the main party entering into this agreement with ACME Corp?".
2.  **Be Conversational:** Be friendly and natural.
3.  **One Question at a Time:** Ask for only ONE piece of information.
4.  **Provide Examples:** After your question, add a new line with a brief, italicized example.
    *E.g., "What is the full name of the client?"*
    *E.g., "Jane Doe"*
5.  **Keep it Brief:** Your questions should be clear and concise.
6.  **Confirmation:** When you get an answer, confirm it briefly ("Got it.", "Perfect.") and then ask the next question.
"""

# --- 4. Helper Functions ---

def extract_text_and_placeholders(file_bytes):
    """
    Extracts text, unique placeholders, and a context map.
    The context map links each placeholder to the paragraph it was found in.
    """
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []
        placeholders = []
        placeholder_context_map = {}

        # Regex to find all common placeholder formats
        placeholder_regex = r"\{{1,2}.*?\}{1,2}|\[.*?\]|<.*?>|%.*?%|__.*?__|\$[a-zA-Z0-9_]+"
        
        def process_paragraphs(paragraphs):
            for para in paragraphs:
                full_text.append(para.text)
                found_in_para = re.findall(placeholder_regex, para.text)
                if found_in_para:
                    placeholders.extend(found_in_para)
                    for ph in found_in_para:
                        if ph not in placeholder_context_map: # Only store first context
                            placeholder_context_map[ph] = para.text

        # Process main body paragraphs
        process_paragraphs(doc.paragraphs)
        
        # Process table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_paragraphs(cell.paragraphs)

        full_text_str = "\n".join(full_text)
        unique_placeholders = sorted(list(set(placeholders)), key=placeholders.index) # Preserve order
        
        return full_text_str, unique_placeholders, placeholder_context_map
    
    except Exception as e:
        st.error(f"Error reading .docx file: {e}")
        return None, [], {}

def replace_placeholders_in_doc(file_bytes, replacements):
    """Replaces placeholders in the docx file and returns the file bytes."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        for p in doc.paragraphs:
            for key, value in replacements.items():
                if key in p.text:
                    for run in p.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, str(value))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in p.text:
                                for run in p.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, str(value))
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.getvalue()
    except Exception as e:
        st.error(f"Error replacing placeholders: {e}")
        return None

def get_ai_response(prompt_to_ai):
    """Sends a prompt to the AI and gets a response."""
    st.session_state.api_history.append({"role": "user", "content": prompt_to_ai})
    
    response = client.chat.completions.create(
        messages=st.session_state.api_history,
        model=COHERE_MODEL
    )
    
    response_text = response.choices[0].message.content
    st.session_state.api_history.append({"role": "assistant", "content": response_text})
    return response_text

def reset_chat():
    """Resets the chat state but keeps the uploaded document."""
    st.session_state.messages = []
    st.session_state.filled_values = {}
    st.session_state.current_placeholder_index = 0
    st.session_state.api_history = [{"role": "system", "content": SYSTEM_PROMPT}]
    
    # Kick off the chat again
    if st.session_state.placeholders:
        try:
            first_ph = st.session_state.placeholders[0]
            context = st.session_state.placeholder_context_map.get(first_ph, "")
            prompt_to_ai = f"Hello! Let's start filling this document. The first placeholder is '{first_ph}'. The context is: \"{context}\". Please ask me the first question."
            
            response_text = get_ai_response(prompt_to_ai)
            st.session_state.messages.append({"role": "assistant", "content": response_text})
        except Exception as e:
            st.error(f"Error with Cohere API: {e}")
    st.rerun()

def clear_full_session_state():
    """Resets the entire session, including the uploaded file."""
    keys_to_clear = list(st.session_state.keys())
    for key in keys_to_clear:
        del st.session_state[key]
    
    # Re-initialize
    initialize_session_state()


# --- 5. Session State Initialization ---
def initialize_session_state():
    """Initializes all required session state variables."""
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "placeholders" not in st.session_state:
        st.session_state.placeholders = []
    if "placeholder_context_map" not in st.session_state:
        st.session_state.placeholder_context_map = {}
    if "filled_values" not in st.session_state:
        st.session_state.filled_values = {}
    if "current_placeholder_index" not in st.session_state:
        st.session_state.current_placeholder_index = 0
    if "original_doc_bytes" not in st.session_state:
        st.session_state.original_doc_bytes = None
    if "original_text" not in st.session_state:
        st.session_state.original_text = ""
    if "api_history" not in st.session_state:
        st.session_state.api_history = [{"role": "system", "content": SYSTEM_PROMPT}]

# Run initialization on first load
initialize_session_state()

# --- 6. Main App UI ---
st.title("✍️ LegalEase AI: Smart Document Filler")
st.markdown("Upload your `.docx` template, and I'll help you fill it in conversationally using AI.")

col1, col2 = st.columns([1, 1], gap="large")

# --- Column 1: Upload & Chat ---
with col1:
    st.header("1. Get Started")
    
    uploaded_file = st.file_uploader(
        "Upload your .docx template", 
        type=["docx"],
        on_change=clear_full_session_state # Reset everything on new file
    )
    
    st.divider()

    # --- File Upload Logic ---
    if uploaded_file is not None and st.session_state.original_doc_bytes is None:
        with st.spinner("Analyzing document..."):
            file_bytes = uploaded_file.getvalue()
            text, placeholders, context_map = extract_text_and_placeholders(file_bytes)
            
            if text is None:
                st.error("Could not read the document.")
            elif not placeholders:
                st.warning("📄 No placeholders (like {Name} or [Date]) found in this document.")
            else:
                st.session_state.original_doc_bytes = file_bytes
                st.session_state.original_text = text
                st.session_state.placeholders = placeholders
                st.session_state.placeholder_context_map = context_map
                st.success(f"Found {len(placeholders)} placeholders!")
                
                # Kick off the chat
                reset_chat() # This will get the first question
                st.rerun() # Rerun to display the first question

    # --- Chat Interface ---
    if st.session_state.original_doc_bytes is not None:
        st.subheader("💬 Chat to Fill")
        
        if st.button("Start Over"):
            reset_chat()
            
        # UI FIX: Put chat history in a container
        chat_container = st.container(height=400, border=True)
        with chat_container:
            for message in st.session_state.messages:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])

        # UI FIX: Chat input is now at the bottom
        if prompt := st.chat_input("Your answer..."):
            # Display user message
            st.session_state.messages.append({"role": "user", "content": prompt})
            with chat_container.chat_message("user"):
                st.markdown(prompt)

            # Process the answer
            with st.spinner("Thinking..."):
                try:
                    # Store the value
                    current_index = st.session_state.current_placeholder_index
                    current_ph = st.session_state.placeholders[current_index]
                    st.session_state.filled_values[current_ph] = prompt
                    
                    # Move to next placeholder
                    st.session_state.current_placeholder_index += 1
                    next_index = st.session_state.current_placeholder_index
                    
                    if next_index < len(st.session_state.placeholders):
                        # --- Ask next question ---
                        next_ph = st.session_state.placeholders[next_index]
                        context = st.session_state.placeholder_context_map.get(next_ph, "")
                        ai_prompt = f"Great, I've saved that. The next placeholder is '{next_ph}'. The context is: \"{context}\". Please ask me the question for this one."
                    else:
                        # --- All done ---
                        ai_prompt = "That was the last placeholder! Please provide a brief, friendly message letting me know I'm all done and can review the document on the right."

                    # Get and display AI response
                    response_text = get_ai_response(ai_prompt)
                    st.session_state.messages.append({"role": "assistant", "content": response_text})
                    with chat_container.chat_message("assistant"):
                        st.markdown(response_text)
                    
                    # Rerun to update the "Edit" and "Preview" sections instantly
                    st.rerun()

                except Exception as e:
                    st.error(f"Error with Cohere API: {e}")
                    st.session_state.current_placeholder_index -= 1 # Roll back on error

# --- Column 2: Review & Download ---
with col2:
    st.header("2. Review & Download")

    if st.session_state.original_doc_bytes is None:
        st.info("Upload a document on the left to get started.")
    else:
        # --- Progress Bar ---
        progress = 0
        if st.session_state.placeholders:
            progress = len(st.session_state.filled_values) / len(st.session_state.placeholders)
        
        st.progress(progress, text=f"{len(st.session_state.filled_values)} / {len(st.session_state.placeholders)} fields filled")
        
        all_filled = len(st.session_state.filled_values) == len(st.session_state.placeholders)
        
        st.divider()

        # --- NEW: Edit Filled Values ---
        with st.expander("✍️ Edit Filled Values"):
            if not st.session_state.filled_values:
                st.write("No values filled yet.")
            else:
                for ph in st.session_state.placeholders:
                    if ph in st.session_state.filled_values:
                        # Use text_input to allow edits
                        st.session_state.filled_values[ph] = st.text_input(
                            f"**{ph}**", 
                            value=st.session_state.filled_values[ph],
                            key=f"edit_{ph}"
                        )
        
        # --- Live Preview ---
        st.subheader("Live Preview")
        with st.container(height=300, border=True):
            preview_text = st.session_state.original_text
            for ph, val in st.session_state.filled_values.items():
                preview_text = preview_text.replace(ph, f"**{val}**") # Bold filled values
            for ph in st.session_state.placeholders:
                if ph not in st.session_state.filled_values:
                    preview_text = preview_text.replace(ph, f"_{ph}_") # Italicize unfilled
            st.markdown(preview_text)
        
        st.divider()
        
        # --- Download Button ---
        if all_filled:
            st.success("All fields filled! 🎉")
            with st.spinner("Generating final document..."):
                final_doc_bytes = replace_placeholders_in_doc(
                    st.session_state.original_doc_bytes,
                    st.session_state.filled_values
                )
            
            if final_doc_bytes:
                st.download_button(
                    label="⬇️ Download Completed Document",
                    data=final_doc_bytes,
                    file_name=f"completed_{uploaded_file.name}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        else:
            st.warning("Please fill in all remaining fields to enable download.")
